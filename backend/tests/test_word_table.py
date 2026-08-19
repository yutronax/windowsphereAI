# Saga #327: word-tablo-basligi (TEST-FIRST / red step) -
# `backend/word_table.py` (append_table) henuz YOK. Bu testler simdilik
# KIRMIZI kalmali (ModuleNotFoundError/ImportError - EXCEL_APPEND'in
# red-step testleriyle AYNI desen).
# Referans: artifacts/word-tablo-basligi/atdd.md (AC-1, AC-2, AC-3, AC-4),
# plan.md (append_table(source_path, headers, rows, backup_path) imzasi,
# excel_rows.append_excel_rows'un backup+tempfile+atomik-replace deseni).

import pytest
from docx import Document

from backend.word_table import append_table


def _write_docx(path, previous_text: str = "mevcut içerik") -> None:
    doc = Document()
    doc.add_paragraph(previous_text)
    doc.save(str(path))


# ---------------------------------------------------------------------------
# append_table(source_path, headers, rows, backup_path)
# ---------------------------------------------------------------------------


def test_append_table_adds_a_table_with_a_header_row_and_keeps_previous_content(tmp_path):
    # AC-1: headers verilince EN ÜSTTE başlık satırı olan bir tablo eklenir,
    # önceki içerik korunur.
    source = tmp_path / "kaynak.docx"
    backup = tmp_path / "yedek" / "kaynak.docx"
    _write_docx(source, "mevcut içerik")

    append_table(
        source,
        headers=["Ad", "Tutar"],
        rows=[["Ali", "100"], ["Veli", "200"]],
        backup_path=backup,
    )

    doc = Document(str(source))
    assert doc.paragraphs[0].text == "mevcut içerik"
    table = doc.tables[-1]
    assert len(table.rows) == 3
    assert len(table.columns) == 2
    assert [cell.text for cell in table.rows[0].cells] == ["Ad", "Tutar"]
    assert [cell.text for cell in table.rows[1].cells] == ["Ali", "100"]
    assert [cell.text for cell in table.rows[2].cells] == ["Veli", "200"]


def test_append_table_without_headers_adds_only_data_rows_and_raises_no_error(tmp_path):
    # AC-2: headers=None -> SADECE veri satırlarından oluşan bir tablo
    # eklenir, hata FIRLATILMAZ.
    source = tmp_path / "kaynak.docx"
    backup = tmp_path / "yedek" / "kaynak.docx"
    _write_docx(source, "mevcut içerik")

    append_table(source, headers=None, rows=[["Ali", "100"]], backup_path=backup)

    doc = Document(str(source))
    assert doc.paragraphs[0].text == "mevcut içerik"
    table = doc.tables[-1]
    assert len(table.rows) == 1
    assert len(table.columns) == 2
    assert [cell.text for cell in table.rows[0].cells] == ["Ali", "100"]


def test_append_table_raises_on_column_count_mismatch_and_leaves_document_unchanged(tmp_path):
    # AC-3: headers 3 sütunlu ama bir rows satırı 2 hücreli -> exception,
    # belgeye HİÇBİR ŞEY eklenmez, kaynak değişmez.
    source = tmp_path / "kaynak.docx"
    backup = tmp_path / "yedek" / "kaynak.docx"
    _write_docx(source, "mevcut içerik")
    pre_bytes = source.read_bytes()

    with pytest.raises(Exception):
        append_table(
            source,
            headers=["Ad", "Tutar", "Şehir"],
            rows=[["Ali", "100"]],
            backup_path=backup,
        )

    assert source.read_bytes() == pre_bytes
    doc = Document(str(source))
    assert len(doc.tables) == 0
    assert len(doc.paragraphs) == 1


def test_append_table_raises_when_source_is_missing(tmp_path):
    # AC-4: kaynak yok -> exception.
    source = tmp_path / "yok.docx"
    backup = tmp_path / "yedek" / "yok.docx"

    with pytest.raises(Exception):
        append_table(source, headers=["Ad"], rows=[["Ali"]], backup_path=backup)

    assert not source.exists()
    assert not backup.exists()


def test_append_table_raises_when_source_is_corrupt_and_leaves_it_untouched(tmp_path):
    # AC-4: kaynak bozuk (python-docx açamıyor) -> exception, kaynak
    # değişmez.
    source = tmp_path / "bozuk.docx"
    corrupt_bytes = b"not a real docx"
    source.write_bytes(corrupt_bytes)
    backup = tmp_path / "yedek" / "bozuk.docx"

    with pytest.raises(Exception):
        append_table(source, headers=["Ad"], rows=[["Ali"]], backup_path=backup)

    assert source.read_bytes() == corrupt_bytes


def test_append_table_backs_up_the_source_before_writing_the_new_content(tmp_path):
    # `backup_path`'in yazmadan ÖNCE alındığını doğrula (backup dosyasının
    # İÇERİĞİ, tablo eklenmeden ÖNCEKİ belgeyle eşleşmeli).
    source = tmp_path / "kaynak.docx"
    backup = tmp_path / "yedek" / "kaynak.docx"
    _write_docx(source, "mevcut içerik")
    pre_append_bytes = source.read_bytes()

    append_table(source, headers=["Ad"], rows=[["Ali"]], backup_path=backup)

    assert backup.exists()
    assert backup.read_bytes() == pre_append_bytes
    backup_doc = Document(str(backup))
    assert backup_doc.paragraphs[0].text == "mevcut içerik"
    assert len(backup_doc.tables) == 0
